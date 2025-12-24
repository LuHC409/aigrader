from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

try:
    from docx import Document  # type: ignore
except ImportError as exc:  # pragma: no cover - dependency injected at runtime
    Document = None  # type: ignore
    _IMPORT_ERROR = exc
else:
    _IMPORT_ERROR = None

from .chunking import estimate_tokens
from .types import DocMeta


class UnsupportedDocumentError(Exception):
    pass


class DocumentExtractionError(Exception):
    pass


def _ensure_docx_available() -> None:
    if Document is None:  # pragma: no cover - handled when dependency missing
        raise RuntimeError(
            "python-docx is required to extract .docx files" + (f": {_IMPORT_ERROR}" if _IMPORT_ERROR else "")
        )


def _clean_lines(lines: List[str]) -> List[str]:
    cleaned: List[str] = []
    last_blank = True
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if not last_blank:
                cleaned.append("")
            last_blank = True
        else:
            cleaned.append(stripped)
            last_blank = False
    return cleaned


def _extract_tables(document) -> List[str]:
    rows: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        rows.append("")
    return rows


def extract_text(path: str, include_tables: bool = True) -> Tuple[str, DocMeta]:
    _ensure_docx_available()
    filepath = Path(path)
    if filepath.suffix.lower() != ".docx":
        raise UnsupportedDocumentError("Only .docx files are supported. Please convert the file before processing.")

    try:
        document = Document(str(filepath))
    except Exception as exc:  # noqa: E722
        raise DocumentExtractionError(f"Failed to read document: {filepath}") from exc

    paragraphs = [p.text for p in document.paragraphs]
    paragraph_count = len(paragraphs)
    lines = _clean_lines(paragraphs)

    table_lines: List[str] = []
    table_count = len(document.tables)
    if include_tables and table_count:
        table_lines = _extract_tables(document)

    combined_lines = lines
    if table_lines:
        combined_lines = lines + [""] + table_lines

    text = "\n".join(combined_lines).strip()
    meta = DocMeta(
        paragraph_count=paragraph_count,
        table_count=table_count,
        char_count=len(text),
    )
    meta.token_est = estimate_tokens(text)
    return text, meta
